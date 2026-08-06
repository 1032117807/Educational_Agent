from __future__ import annotations

import html
from pathlib import Path


def markdown_to_html(markdown: str) -> str:
    lines: list[str] = []
    in_list = False
    for raw_line in markdown.splitlines():
        line = html.escape(raw_line)
        if line.startswith("# "):
            lines.append(f"<h1>{line[2:]}</h1>")
        elif line.startswith("## "):
            lines.append(f"<h2>{line[3:]}</h2>")
        elif line.startswith("- "):
            if not in_list:
                lines.append("<ul>")
                in_list = True
            lines.append(f"<li>{line[2:]}</li>")
        else:
            if in_list:
                lines.append("</ul>")
                in_list = False
            lines.append("<p></p>" if not line else f"<p>{line}</p>")
    if in_list:
        lines.append("</ul>")
    return "<!doctype html><html><head><meta charset='utf-8'><style>body{font-family:Arial,'Microsoft YaHei',sans-serif;margin:42px;line-height:1.6}h1{color:#1b4965}h2{border-bottom:1px solid #ccd6dd;padding-bottom:4px}</style></head><body>" + "\n".join(lines) + "</body></html>"


def export_report(markdown: str, target: Path) -> None:
    suffix = target.suffix.lower()
    if suffix == ".md":
        target.write_text(markdown, encoding="utf-8")
    elif suffix == ".html":
        target.write_text(markdown_to_html(markdown), encoding="utf-8")
    elif suffix == ".docx":
        from docx import Document

        document = Document()
        for line in markdown.splitlines():
            if line.startswith("# "):
                document.add_heading(line[2:], level=1)
            elif line.startswith("## "):
                document.add_heading(line[3:], level=2)
            elif line.startswith("- "):
                document.add_paragraph(line[2:], style="List Bullet")
            elif line:
                document.add_paragraph(line)
        document.save(target)
    elif suffix == ".pdf":
        from PySide6.QtCore import QMarginsF
        from PySide6.QtGui import QPageLayout, QPageSize, QPdfWriter, QTextDocument

        writer = QPdfWriter(str(target))
        writer.setPageLayout(QPageLayout(QPageSize(QPageSize.A4), QPageLayout.Portrait, QMarginsF(18, 18, 18, 18)))
        document = QTextDocument()
        document.setHtml(markdown_to_html(markdown))
        document.print_(writer)
    else:
        raise ValueError("Unsupported report format")
